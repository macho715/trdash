from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\tr_dash-main")
SOURCE = ROOT / "work" / "TR5_Schedule_Final_v3.md"
OUT_DIR = ROOT / "output" / "doc"
OUTPUT = OUT_DIR / "TR5_Schedule_Final_v3_EN_styled.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "Table"
    p.add_run(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header)
        shade_cell(hdr[idx], "D9EAF7")
        if widths:
            hdr[idx].width = Inches(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Inches(widths[idx])

    doc.add_paragraph()


def add_meta(doc: Document, label: str, value: str) -> None:
    doc.add_paragraph(f"{label}: {value}", style="Body")


def add_bullet(doc: Document, text: str, checked: bool | None = None) -> None:
    prefix = ""
    if checked is False:
        prefix = "[ ] "
    elif checked is True:
        prefix = "[x] "
    doc.add_paragraph(f"- {prefix}{text}", style="Body")


def add_note(doc: Document, text: str) -> None:
    doc.add_paragraph(f"Note: {text}", style="Caption")


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        doc.add_paragraph(line, style="Caption")


def ensure_paragraph_style(doc: Document, name: str, size: float, bold: bool = False, font: str = "Aptos") -> None:
    styles = doc.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    ensure_paragraph_style(doc, "Body", 9.5)
    ensure_paragraph_style(doc, "Heading1", 14, bold=True)
    ensure_paragraph_style(doc, "Heading2", 12, bold=True)
    ensure_paragraph_style(doc, "Heading3", 10.5, bold=True)
    ensure_paragraph_style(doc, "Caption", 8, font="Consolas")
    ensure_paragraph_style(doc, "Table", 8)
    doc.styles["Heading1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def heading(doc: Document, text: str, level: int) -> None:
    style = "Heading1" if level <= 1 else "Heading2" if level == 2 else "Heading3"
    doc.add_paragraph(text, style=style)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)

    heading(doc, "TR5 (5th Voyage) Pre-Operation Schedule - Final Plan", 1)

    add_meta(doc, "Document ID", "HVDC-TR5-PREOP-FINAL-001")
    add_meta(doc, "Revision", "v8.0")
    add_meta(doc, "Issued", "2026-03-29 (Asia/Dubai)")
    add_meta(doc, "T+0 (Go Signal)", "2026-04-26 (Saturday)")
    add_meta(doc, "LCT Bushra MZP All Fast", "2026-05-04 (Sunday, T+8)")
    add_meta(doc, "Target RoRo Load-out", "2026-05-06 (Tuesday, T+10)")
    add_meta(doc, "Target AGI Jacking Down Complete", "2026-05-13 (Tuesday, T+17)")
    add_meta(doc, "Critical Path Float", "0.00d (TR5 to TR6 to TR7 serial sequence)")

    heading(doc, "1. TR1-TR4 Actual History Summary", 1)
    heading(doc, "1.1 Key Actual Results by Voyage", 2)
    add_table(
        doc,
        ["Voyage", "MZP All Fast", "ETD", "AGI Berthing", "Port Turn", "Remarks"],
        [
            ["TR1", "2026-01-27 17:38", "2026-01-31 20:36", "2026-02-04 16:06", "3.17d", "Completed 2d earlier than plan"],
            ["TR2", "2026-02-09 21:18", "2026-02-12 15:00", "2026-02-13-14", "2.74d", "Dense fog on Feb 11 caused +3d delay"],
            ["TR3", "2026-02-17 14:24", "2026-02-19 03:00", "2026-02-20-21", "1.52d", "G27kt strong wind on Feb 16 caused +1d delay"],
            ["TR4", "2026-02-26 (est.)", "2026-02-27 08:42", "2026-02-28 07:12", "1.6d", "Feb 20 anchorage -> approx. Feb 22 berthing -> Feb 24 Load-out -> Feb 27 08:42 ETD actual"],
        ],
        [0.6, 1.1, 1.1, 1.1, 0.7, 2.8],
    )

    heading(doc, "1.2 Confirmed SSOT v1.1 Segments", 2)
    add_table(
        doc,
        ["Segment", "Confirmed Value", "Basis"],
        [
            ["Port Turn (MZP All Fast to ETD)", "3.00d", "TR4: Feb 26 to Mar 1"],
            ["Sea Leg (ETD to AGI Berthing)", "1.00d", "TR4: Mar 1 to Mar 2"],
            ["AGI Berthing to Jacking Down Complete", "6.00d", "TR4: Mar 2 to Mar 8"],
        ],
        [2.4, 1.3, 3.0],
    )
    add_note(
        doc,
        "TR5 Port Turn plan is 2d (May 4 All Fast -> May 6 Load-out) versus the TR1-TR4 average of 2.6d. "
        "The actual TR1-TR4 range is 1.5d (TR3 shortest) to 4.5d (TR1 delayed). "
        "A 2d turn is close to TR3 performance and requires all documents to be completed in advance, PTW G1 by Apr 29, HM G2 by May 4, and immediate Linkspan installation on May 5. "
        "If even one document package is not ready, Port Turn may extend by +1-2d, moving ETD to May 7-8 and JD Complete to May 14-15."
    )

    heading(doc, "1.2-B Actual Document Lead Times for TR1-TR4 (WhatsApp Basis)", 2)
    add_note(doc, "Basis: All Fast = D+0. Negative values mean before All Fast; positive values mean after All Fast.")
    add_table(
        doc,
        ["Voyage", "PTW Submit", "PTW Approval", "HM Contact", "Linkspan", "Pre-Op Meeting", "Load-out", "Port Turn"],
        [
            ["TR1", "D+0 (Jan27)", "D+3 (Jan30 PM exception)", "D+1 F2F (Jan28)", "D+0 attempt -> D+4 exception", "D+2 F2F", "D+4 (Jan31)", "4.1d"],
            ["TR2", "D-4 (Feb5)", "D+0 (Feb9 AM)", "D+0 pre-op", "D+2 plan -> D+3 (fog)", "D+0 (Feb9 11:00)", "D+3 (Feb12)", "2.7d"],
            ["TR3", "D-2 date change", "D+0 same day", "D+0 continuous", "D+0 loaded on LCT", "D+0 (Feb17 15:00)", "D+1 (Feb18)", "1.5d"],
            ["TR4", "D-1 (Feb23)", "D+0 (Feb24 AM)", "D-1 15:00", "D-1", "D-4 Port FRA", "D+1-2", "~2.0d"],
            ["Average", "D-2", "D+0-1", "D+0-1", "D+0-1", "D+0", "D+1-3", "~2.6d"],
        ],
        [0.6, 0.9, 1.0, 0.85, 1.0, 1.0, 0.9, 0.7],
    )
    add_note(
        doc,
        "Why the TR5 preparation period is extended to 8 days compared with TR1-TR4: "
        "1) a new 2nd SPMT mobilization requires 8 days; 2) LCT Bushra repositioning requires an 8-day voyage; "
        "3) restart after a two-month gap means all documents have expired and require 8-10 days for reissue; "
        "4) the approval timing itself for G1 and G2 remains aligned with the TR2-TR4 average pattern of D+0-1."
    )

    heading(doc, "1.3 Lessons Learned from TR1-TR4 (WhatsApp Actual Basis)", 2)
    add_note(doc, "Actual TR1-TR4 Port Turn range: 1.5d (TR3 shortest) to 5d (including TR4 anchorage).")
    add_table(
        doc,
        ["#", "Actual Event (WhatsApp)", "TR5 Application"],
        [
            ["L-1", "TR1: PTW was not approved until Berth Day 3. SPMT movement stopped completely and Load-out proceeded only after HM exception approval.", "PTW application at T+0 (Apr 26) is mandatory. Escalate immediately if not approved before G1 (Apr 29)."],
            ["L-2", "TR1: Linkspan certificate was not ready. HM did not permit loading and allowed exception only after a formal meeting. Reconfirm this requirement for every voyage.", "Pkg-D (Linkspan Cert) must be submitted to HM at T+1 (Apr 27)."],
            ["L-3", "TR2: dense fog on Feb 11 delayed Load-out by 1 day (Feb 12 ETD).", "Hold if G4 visibility is below 2 NM. Secure May 7 buffer."],
            ["L-4", "TR3: G27kt strong wind caused a 1-day delay before Feb 17 arrival and Feb 19 ETD.", "Maintain G4 wind criterion of 15 kt or below."],
            ["L-5", "TR4: Berth #5 occupancy caused anchorage waiting from Feb 20 to approx. Feb 22, resulting in 2d lost.", "Apply for Berth Booking immediately at T+0 and secure Jetty #5."],
            ["L-6", "Current SPMT is on AGI Island and cannot be brought back to MZP in time.", "Mobilize 2nd SPMT under Option C. Issue Mammoet order immediately at T+0."],
            ["L-7", "Documents expired for every voyage and all packages required new issuance.", "Confirm with HM whether existing TR1 engineering documents can be reused."],
            ["L-8", "TR2: FRA and Pre-op meeting were completed on Berth Day 1.", "Maintain sequence: FRA on Berth Day 1 (May 5), then Pre-op and Load-out on Day 2 (May 6)."],
        ],
        [0.4, 3.2, 3.2],
    )

    heading(doc, "2. TR5 Operating Conditions", 1)
    heading(doc, "2.1 SPMT Status", 2)
    add_bullet(doc, "Current status: one SPMT is located on AGI Island and cannot be brought back to MZP.")
    add_bullet(doc, "Decision: Option C - Mammoet to mobilize a new 2nd SPMT directly to MZP.")
    add_bullet(doc, "Required duration: order at T+0 (Apr 26) -> arrival at MZP by T+8 (May 4).")

    heading(doc, "2.2 Linkspan Installation Conditions", 2)
    add_bullet(doc, "Mandatory requirement: install during Deck Preparation before Load-out.")
    add_bullet(doc, "Installation date: Berth Day 1 (May 5, T+9), in parallel with Deck Prep.")
    add_bullet(doc, "Preparation: transport order on T+0, arrival at MZP on T+3 (Apr 29), installation on LCT after Bushra berths on May 4.")

    heading(doc, "2.3 New Permit and Document Reissue", 2)
    add_note(doc, "Premise: all documents used through TR4 have expired and must be newly issued.")
    add_table(
        doc,
        ["Package", "Description", "Application Timing", "Lead Time", "TR Actual Average"],
        [
            ["Pkg-A Customs Declaration", "Maqta customs declaration", "T+0", "3 WD", "3 WD, same for every voyage"],
            ["Pkg-B HM Mooring + SPMT Stowage", "Harbour Master review package", "T+1", "6 WD (UAE Sun-Thu)", "TR2: D-4 submit -> D+0 approval. TR5 has 8d float due to SPMT/restart."],
            ["Pkg-D Linkspan Cert Drawing MOC", "Vessel ramp certification", "T+1", "5 days", "TR2: D-4 submit -> D+2-3 use. TR3: loaded on board. TR4: D-1."],
            ["Pkg-E PTW", "Hot Work / WOW / SPMT Land permits", "T+0", "1-2 WD", "TR2: D-4 submit -> D+0 receipt. TR4: D-1 apply -> D+0 receipt. Actual 1-2 WD."],
            ["Pkg-F Berth Booking", "AD Ports berth reservation", "T+0", "2 days", "TR4: Feb23 HM -> Feb24 Load-out. Gate pass can be processed same day."],
            ["Pkg-G Ramp Certificate Voyage 5", "KFS voyage certification", "T+0", "7 days", "Required every voyage. Standard 7-day lead time."],
        ],
        [1.3, 1.6, 1.1, 0.8, 2.4],
    )

    heading(doc, "3. TR5 Phase-Gate Schedule", 1)
    heading(doc, "3.1 Gantt Chart Source", 2)
    add_code_block(
        doc,
        """
gantt
    title TR5 Pre-Op T+0=Apr26 Bushra=May04 Load-out=May06 JD=May13
    dateFormat YYYY-MM-DD
    axisFormat %b-%d

    section OFCO Customs and PTW
    Customs Declaration Maqta PkgA      :doc1, 2026-04-26, 3d
    PTW Submit HotWork WOW SPMT Land    :crit, doc2, 2026-04-26, 2d
    PTW Approval 1to2 Working Days      :crit, doc3, 2026-04-27, 2d
    Berth Booking AD Ports Request      :brq, 2026-04-26, 2d
    G1 PTW Approved                     :crit, milestone, g1, 2026-04-29, 0d

    section OFCO MMT Harbour Master Approval
    HM PkgB Mooring SPMT Stowage Plan   :hm1, 2026-04-27, 1d
    HM PkgD Linkspan Cert Drawing MOC   :hm2, 2026-04-27, 1d
    HM Review Period CRITICAL PATH      :crit, hm3, 2026-04-27, 7d
    G2 HM Approval Received             :crit, milestone, g2, 2026-05-04, 0d

    section Mammoet Equipment Mobilization
    SPMT 2nd Order and Transit to MZP   :eq1, 2026-04-26, 8d
    SPMT Arrives MZP Inspection         :eq2, 2026-05-04, 1d
    Linkspan Transport to MZP           :eq3, 2026-04-26, 3d
    Linkspan LCT Install with Deck Prep :crit, eq4, 2026-05-05, 1d
    G2-EQ All Equipment Ready at MZP    :milestone, geq, 2026-05-05, 0d

    section KFS LCT Bushra
    Bushra Repositioning to MZP         :lct1, 2026-04-26, 8d
    Ramp Certificate Voyage 5           :crit, lct2, 2026-04-26, 7d
    Bushra All Fast MZP Berth Secured   :milestone, bf, 2026-05-04, 0d

    section ALL MZP Berth Ops
    Berth Day 1 Deck Prep Stool Weld    :dp1, 2026-05-05, 1d
    Berth Day 1 Linkspan Install        :crit, li1, 2026-05-05, 1d
    Berth Day 1 Beam Change FW Supply   :crit, bd1, 2026-05-05, 1d
    Berth Day 1 FRA Formal Risk         :fra1, 2026-05-05, 1d
    Berth Day 2 FRA Port Ops Morning    :crit, fra2, 2026-05-06, 1d
    Berth Day 2 Pre-Op Meeting          :pop, 2026-05-06, 1d
    Berth Day 2 RoRo Load-out           :crit, bd2, 2026-05-06, 1d
    Berth Day 2 Sea Fastening Complete  :crit, sf, 2026-05-06, 1d
    G4 Weather and Tide Go-No-Go        :crit, milestone, g4, 2026-05-06, 0d
    Berth Day 3 Weather Buffer          :buf, 2026-05-07, 1d

    section KFS MMT Voyage and AGI Ops
    Voyage MZP to AGI Sea Leg 1d        :crit, voy, 2026-05-06, 1d
    AGI All Fast Berthing Confirmed     :milestone, ab, 2026-05-07, 0d
    AGI Load-in Offloading Complete     :li, 2026-05-07, 1d
    AGI JD Transport to Laydown Yard    :jdt, 2026-05-08, 2d
    AGI Jacking Down 6d Cycle           :crit, jd, 2026-05-07, 6d
    TARGET AGI Jacking Down Complete    :crit, milestone, jdc, 2026-05-13, 0d
        """,
    )

    doc.add_page_break()
    heading(doc, "3.2 Detailed Timeline (T+0 to T+17)", 2)
    add_table(
        doc,
        ["Day", "Date", "Weekday", "Activity", "Owner", "Gate"],
        [
            ["T+0", "Apr 26", "Sat (weekend)", "Go Signal. Kick-off meeting by phone/email. Apply for three PTWs. Submit Customs. Order Linkspan transport. Order 2nd SPMT. Instruct Bushra positioning. Apply for Ramp Cert.", "Samsung/OFCO/MMT/KFS", "-"],
            ["T+1", "Apr 27", "Sun", "Officially submit HM Pkg-B and Pkg-D. HM Review starts. PTW officially received. Apply for Berth Booking.", "OFCO/MMT", "-"],
            ["T+2", "Apr 28", "Mon", "PTW under review. Confirm existing engineering documents. Bushra and SPMT in transit.", "-", "-"],
            ["T+3", "Apr 29", "Tue", "G1: three PTWs approved. Linkspan arrives at MZP. Customs completed.", "OFCO", "G1 PTW"],
            ["T+4", "Apr 30", "Wed", "HM review continues. Bushra and SPMT in transit.", "-", "-"],
            ["T+5", "May 1", "Thu", "HM review continues. Ramp Cert in progress.", "-", "-"],
            ["T+6", "May 2", "Fri (weekend)", "-", "-", "-"],
            ["T+7", "May 3", "Sat (weekend)", "HM review completed (6 calendar days: Apr 27-May 3).", "-", "-"],
            ["T+8", "May 4", "Sun", "Bushra MZP All Fast. 2nd SPMT arrives at MZP. Equipment inspection. G2 HM Approval issued. Ramp Cert completed.", "All", "G2 HM + G2-EQ"],
            ["T+9", "May 5", "Mon", "Berth Day 1: Deck Prep, Linkspan installation in parallel, Beam Change, FW Supply, FRA, and 2nd SPMT setup on board.", "All", "-"],
            ["T+10", "May 6", "Tue", "Berth Day 2 - TARGET LOAD-OUT. Pre-Op Meeting/Toolbox Talk. FRA Port Ops. RoRo Load-out by SPMT during 11:00-14:00 tide window. Sea Fastening. Tide/Weather Go-No-Go. ETD.", "All", "G4 Tide"],
            ["T+11", "May 7", "Wed", "Buffer for weather hold or AGI Berthing (All Fast). Start AGI Load-in.", "KFS/ADNOC", "-"],
            ["T+11-12", "May 7-8", "-", "Complete AGI Load-in Offloading. Start AGI JD Transport.", "MMT", "-"],
            ["T+11-17", "May 7-13", "-", "AGI Jacking Down 6.00d Cycle.", "MMT/AGI", "-"],
            ["T+17", "May 13", "Tue", "TARGET: AGI Jacking Down Complete.", "MMT", "FINAL"],
        ],
        [0.6, 0.65, 0.8, 3.6, 1.0, 0.9],
    )
    add_note(doc, "UAE weekend note: Friday/Saturday are weekends. Marine and Mammoet equipment movements continue on Apr 26, May 2, and May 3, while official government agency work operates Sunday through Thursday.")

    heading(doc, "4. Gate Conditions", 1)
    heading(doc, "G1 - PTW Approved (Apr 29, T+3)", 2)
    for item in [
        "Issuing authority: AD Ports Authority.",
        "Three permits: Hot Work / WOW / SPMT Land PTW.",
        "Application route: T+0 (Apr 26) email -> T+1 (Apr 27) official receipt -> 1-2 WD (Apr 27+28) -> G1 on Apr 29.",
        "Actual basis: TR2 D-4 application -> D+0 receipt (1 WD), TR4 D-1 application -> D+0 receipt (1 WD), TR1 D+0 -> D+3 under crisis response. Average 1-2 WD.",
        "Condition: SPMT movement cannot start without SPMT Land PTW.",
    ]:
        add_bullet(doc, item)

    heading(doc, "G2 - HM Approval (May 4, T+8)", 2)
    for item in [
        "Issuing authority: Harbour Master, Mina Zayed.",
        "Package: Pkg-B (Mooring/SPMT Stowage) + Pkg-D (Linkspan Cert Drawing MOC).",
        "Review: submit on T+1 (Apr 27, UAE WD1) -> 6 WD (Apr 27/28/29/30/May 1/May 4) -> issue on May 4 (Sun).",
        "TR1 lesson: HM requires the Linkspan certificate. If Pkg-D is not submitted, G2 will be rejected.",
        "Key point: any T+1 submission delay directly affects Load-out because G2 may not be ready when Bushra arrives.",
    ]:
        add_bullet(doc, item)

    heading(doc, "G4 - Tide / Weather Go (May 6, T+10 or Buffer May 7)", 2)
    add_table(
        doc,
        ["Item", "Criterion", "Basis"],
        [
            ["Wind Speed", "15 kt or below", "TR3 G27kt lesson"],
            ["Wave Height (Hs)", "1.25 m or below", "Ramp operation limit"],
            ["Visibility", "2 NM or above", "TR2 fog lesson"],
            ["Tide (MWL)", "1.5 m or above", "Ramp draft condition"],
            ["Tidal Window", "11:00-14:00 GST", "May 6 or May 7"],
        ],
        [1.5, 1.6, 2.8],
    )

    heading(doc, "5. Critical Path", 1)
    add_code_block(
        doc,
        """
T+0 Apr26: Go Signal -> PTW application / SPMT order / Bushra departure
        |
T+1 Apr27: HM Pkg-B+D submission -> HM Review 6d starts
        |
T+3 Apr29: G1 PTW Approved
        |
T+8 May04: Bushra All Fast + G2 HM Approval + SPMT arrival
        |
T+9 May05: Berth Day 1 - Deck Prep + Linkspan installation + Beam Change
        |
T+10 May06: Berth Day 2 - RoRo Load-out + G4 Tide GO + ETD
        |
T+11 May07: Voyage 1.00d -> AGI Berthing
        |
T+17 May13: AGI Jacking Down Complete (6.00d)
        """,
    )
    add_bullet(doc, "Float = 0.00d (TR5 -> TR6 -> TR7 serial sequence).")

    heading(doc, "6. Equipment Mobilization Plan", 1)
    add_table(
        doc,
        ["Equipment", "Current Location", "Action", "Duration", "MZP Arrival"],
        [
            ["SPMT 1 unit", "AGI Island", "Keep at current location; recovery not required", "-", "-"],
            ["SPMT 2nd (new)", "Mammoet warehouse", "Order at T+0 (Apr 26) -> 8 days", "8 days", "May 4 (T+8)"],
            ["Linkspan", "AGI or near MZP", "Transport order at T+0 -> 3 days", "3 days", "Apr 29 (T+3)"],
            ["LCT Bushra", "Current location", "Positioning instruction at T+0 -> 8-day voyage", "8 days", "May 4 (T+8)"],
            ["Beam/Cradle", "Previous voyage location", "Transport together with Bushra", "8 days", "May 4 (T+8)"],
        ],
        [1.2, 1.3, 2.4, 0.8, 1.1],
    )
    add_note(doc, "Linkspan installation sequence: (1) arrive at MZP on Apr 29 and store; (2) Bushra All Fast on May 4; (3) install on LCT during Berth Day 1 on May 5 in parallel with Deck Prep to satisfy G2-EQ.")

    heading(doc, "7. Document Package Checklist (All Newly Issued)", 1)
    checklist = {
        "Pkg-A: Customs / OFCO": [
            "Customs Declaration Maqta (item declaration)",
            "Packing List / Commercial Invoice / Bill of Lading",
        ],
        "Pkg-B: Harbour Master": [
            "Mooring Plan (updated for Voyage 5)",
            "SPMT Stowage Plan (reflecting 2nd SPMT)",
            "Ballast Plan / Stability Calculation (HM confirmation required for TR1-TR4 reuse)",
        ],
        "Pkg-C: Engineering (Existing Reuse)": [
            "Stability Report / FEA / Ballast Calculation (existing TR1-TR4 documents; HM reuse approval required)",
        ],
        "Pkg-D: Linkspan Certificate": [
            "Linkspan Drawing MOC / Load Test Certificate / Drawing for HM submission",
        ],
        "Pkg-E: PTW (3 Types)": [
            "Hot Work PTW / Working Over Water (WOW) PTW / SPMT Land PTW",
        ],
        "Pkg-F: Berth/Port": [
            "Berth Booking (AD Ports) / Port Entry Permit / Gate Pass for equipment and vehicles",
        ],
        "Pkg-G: KFS / LCT": [
            "Ramp Certificate Voyage 5 / LCT Crew Visa renewal / KFS Vessel Insurance",
        ],
        "Pkg-H: CICPA (AGI Island Access)": [
            "CICPA renewal - Mammoet work crew AGI Island access permit (AD Ports/CICPA)",
            "KFS crew AGI berthing permit for vessel berthing at AGI Berth",
            "Application timing: immediately at T+0 (Apr 26). Lead time: 5-7 WD. Deadline: before AGI arrival at T+11 (May 7).",
            "Basis: TR2 WhatsApp - AGI movement was delayed by 1 day because Mammoet crew CICPA was not renewed.",
        ],
    }
    for section_heading, items in checklist.items():
        heading(doc, section_heading, 2)
        for item in items:
            add_bullet(doc, item, checked=False)

    heading(doc, "8. Risk Register", 1)
    add_table(
        doc,
        ["#", "Risk", "Probability", "Impact", "Response"],
        [
            ["R-1", "HM Approval delay", "HIGH", "+3d", "Submit at T+1 (Apr 27) without fail. Prepare Pkg-B in advance."],
            ["R-2", "Weather no-go (fog/strong wind)", "MED", "+1-2d", "Secure G4 buffer on May 7."],
            ["R-3", "Bushra voyage delay (weather/engine)", "MED", "+1-2d", "Monitor within 8-day float. Escalate if ETA is not confirmed by T+6."],
            ["R-4", "Linkspan certification delay", "MED", "+2d", "Start immediately at T+0. Submit Pkg-D at T+1. TR1 lesson: HM requires it."],
            ["R-5", "2nd SPMT unavailable", "MED", "+5d", "Confirm with Mammoet immediately at T+0. Backup: recover Bushra from AGI -> +3-4d."],
            ["R-6", "PTW delay due to public holiday or review delay", "MED", "+1-2d", "TR1 lesson: PTW not approved after berthing stopped all operations. Apply immediately at T+0."],
            ["R-7", "MZP Berth occupied", "MED", "+1-4d", "Apply for Berth Booking immediately at T+0. TR4 lesson: Feb 20 anchorage caused 4d waiting."],
            ["R-8", "G2 HM not issued by May 4", "LOW", "+1d", "Complete by May 3 (Sat) and confirm final receipt on May 4 (Sun)."],
            ["R-9", "Tidal Window missed", "MED", "+1d", "Make Go-NoGo decision at 06:00 on May 6."],
            ["R-10", "Iran situation worsens again", "EXT", "Indefinite", "Hold. Wait for ADNOC/Samsung HSE decision."],
            ["R-11", "CICPA renewal delay causing AGI crew access failure", "MED", "+1-2d", "Apply immediately at T+0. TR2 lesson: if not renewed, crew cannot board on AGI arrival day. Complete by T+8 (May 4)."],
        ],
        [0.5, 1.6, 0.8, 0.8, 3.2],
    )

    heading(doc, "9. Action Items by Organization", 1)
    actions = {
        "Samsung C&T (PM)": [
            "T+0 (Apr 26): call kick-off meeting by phone/video.",
            "T+0: formally document Go Signal.",
            "T+0: request OFCO to confirm whether existing engineering documents can be reused.",
            "T+8 (May 4): confirm receipt of HM Approval and instruct Berth Day 1 mobilization.",
        ],
        "Mammoet (MMT)": [
            "T+0 (Apr 26): apply for CICPA renewal for Mammoet crew AGI access (5-7 WD, deadline T+8 May 4).",
            "T+0 (Apr 26): immediately confirm 2nd SPMT availability and place order.",
            "T+0: confirm Linkspan current location and transport lead time.",
            "T+1 (Apr 27): submit HM Pkg-B (SPMT Stowage Plan).",
            "T+9 (May 5): Berth Day 1 - complete SPMT setup, Linkspan installation, and Beam Change.",
        ],
        "OFCO (Documents / PTW)": [
            "T+0 (Apr 26): apply by email for three PTWs.",
            "T+0: submit Customs Maqta Declaration.",
            "T+0: apply for Berth Booking with AD Ports.",
            "T+1 (Apr 27): formally submit HM Pkg-B/D.",
        ],
        "KFS (LCT Bushra)": [
            "T+0 (Apr 26): instruct Bushra to depart for MZP positioning.",
            "T+0: apply for Ramp Certificate Voyage 5.",
            "T+0: confirm crew visa and Gate Pass renewals.",
            "T+8 (May 4): confirm and report MZP All Fast.",
        ],
    }
    for section_heading, items in actions.items():
        heading(doc, section_heading, 2)
        for item in items:
            add_bullet(doc, item, checked=False)

    heading(doc, "10. Key Date Summary", 1)
    add_table(
        doc,
        ["Event", "Date", "Weekday", "Day"],
        [
            ["Go Signal", "2026-04-26", "Sat (weekend)", "T+0"],
            ["HM Pkg-B/D submission", "2026-04-27", "Sun", "T+1"],
            ["G1 PTW Approved", "2026-04-29", "Tue", "T+3"],
            ["Linkspan arrival at MZP", "2026-04-29", "Tue", "T+3"],
            ["Ramp Certificate complete", "2026-05-03", "Sat", "T+7"],
            ["Bushra MZP All Fast", "2026-05-04", "Sun", "T+8"],
            ["G2 HM Approval", "2026-05-04", "Sun", "T+8"],
            ["SPMT 2nd arrival at MZP", "2026-05-04", "Sun", "T+8"],
            ["Berth Day 1 (Deck Prep + Linkspan installation + Beam Change)", "2026-05-05", "Mon", "T+9"],
            ["TARGET RoRo Load-out (Tidal 11:00-14:00 GST)", "2026-05-06", "Tue", "T+10"],
            ["G4 Tide/Weather Go", "2026-05-06", "Tue", "T+10"],
            ["Buffer if Tide NG", "2026-05-07", "Wed", "T+11"],
            ["ETD (voyage start)", "2026-05-06-07", "Tue-Wed", "T+10-11"],
            ["AGI Berthing (All Fast)", "2026-05-07", "Wed", "T+11"],
            ["AGI Load-in complete", "2026-05-07-08", "Wed-Thu", "T+11-12"],
            ["AGI Jacking Down complete (TARGET)", "2026-05-13", "Tue", "T+17"],
        ],
        [3.5, 1.1, 1.0, 0.8],
    )
    add_bullet(doc, "Total duration: T+0 (Apr 26) -> Load-out (May 6) = 10 days.")
    add_bullet(doc, "Total duration: T+0 (Apr 26) -> JD Complete (May 13) = 17 days.")

    heading(doc, "11. Assumptions", 1)
    for item in [
        "T+0 = Apr 26 (Sat/weekend): Go Signal and email/phone orders proceed during the weekend. Official government agency work starts from Apr 27 (Sun).",
        "Engineering document reuse: HM is assumed to allow reuse of existing TR1-TR4 Stability/FEA documents. If rejected, add +5d.",
        "SPMT 2nd availability: assumed to arrive at MZP within 8 days. If unavailable, Bushra AGI recovery under Option A adds +3-4d.",
        "Linkspan location: assumed to be at MZP or nearby and able to arrive within 3 days.",
        "Tide window: at least one acceptable 11:00-14:00 GST tide window is assumed on May 6 or May 7.",
        "Iran situation: Go Signal assumes easing of US-Iran tensions and completion of UAE/ADNOC safety confirmation.",
    ]:
        add_bullet(doc, item)

    add_note(doc, "Document end - HVDC-TR5-PREOP-FINAL-001 Rev v8.0")
    add_note(doc, "Bushra MZP All Fast: 2026-05-04 | Target Load-out: 2026-05-06 | Target AGI JD: 2026-05-13")

    doc.core_properties.title = "TR5 Pre-Operation Schedule - Final Plan"
    doc.core_properties.subject = "English version converted from TR5_Schedule_Final_v3.md"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

