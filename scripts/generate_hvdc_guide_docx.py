"""
Generate HVDC_TR_Standard_Operating_Guide_v1.0.docx from the Markdown source.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

SRC = r"c:\tr_dash-main\docs\HVDC_TR_Standard_Operating_Guide_v1.0.md"
OUT = r"c:\tr_dash-main\docs\HVDC_TR_Standard_Operating_Guide_v1.0.docx"

doc = Document()

# ---- page margins ----
for sec in doc.sections:
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.2)
    sec.right_margin = Inches(1.2)

# ---- styles ----
styles = doc.styles

def set_heading_style(style_name, pt_size, bold=True, color=None):
    if style_name in [s.name for s in styles]:
        s = styles[style_name]
        s.font.size = Pt(pt_size)
        s.font.bold = bold
        if color:
            s.font.color.rgb = RGBColor(*color)

set_heading_style("Heading 1", 16, bold=True, color=(0, 70, 127))
set_heading_style("Heading 2", 13, bold=True, color=(0, 112, 192))
set_heading_style("Heading 3", 11, bold=True, color=(50, 50, 50))

# ---- helper: add styled paragraph ----
def add_para(doc, text, style="Normal", bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# ---- parse markdown table ----
def parse_md_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line or set(line.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue  # separator row
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows

def add_md_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    # normalize
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:  # header
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()

# ---- parse and render ----
with open(SRC, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
table_buffer = []
in_code = False
code_buffer = []

while i < len(lines):
    raw = lines[i].rstrip("\n")
    stripped = raw.strip()

    # Code block toggle
    if stripped.startswith("```"):
        if not in_code:
            in_code = True
            code_buffer = []
        else:
            in_code = False
            code_text = "\n".join(code_buffer)
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            doc.add_paragraph()
        i += 1
        continue

    if in_code:
        code_buffer.append(raw)
        i += 1
        continue

    # Heading 1
    if stripped.startswith("# ") and not stripped.startswith("## "):
        # flush table
        if table_buffer:
            add_md_table(doc, parse_md_table(table_buffer))
            table_buffer = []
        text = stripped[2:].strip()
        doc.add_heading(text, level=1)
        i += 1
        continue

    # Heading 2
    if stripped.startswith("## ") and not stripped.startswith("### "):
        if table_buffer:
            add_md_table(doc, parse_md_table(table_buffer))
            table_buffer = []
        text = stripped[3:].strip()
        doc.add_heading(text, level=2)
        i += 1
        continue

    # Heading 3
    if stripped.startswith("### ") and not stripped.startswith("#### "):
        if table_buffer:
            add_md_table(doc, parse_md_table(table_buffer))
            table_buffer = []
        text = stripped[4:].strip()
        doc.add_heading(text, level=3)
        i += 1
        continue

    # Heading 4
    if stripped.startswith("#### "):
        if table_buffer:
            add_md_table(doc, parse_md_table(table_buffer))
            table_buffer = []
        text = stripped[5:].strip()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        i += 1
        continue

    # HR
    if stripped.startswith("---"):
        if table_buffer:
            add_md_table(doc, parse_md_table(table_buffer))
            table_buffer = []
        doc.add_paragraph()
        i += 1
        continue

    # Table row detection
    if stripped.startswith("|"):
        table_buffer.append(stripped)
        i += 1
        continue
    else:
        if table_buffer:
            add_md_table(doc, parse_md_table(table_buffer))
            table_buffer = []

    # Bullet list
    if stripped.startswith("- ") or stripped.startswith("* "):
        text = stripped[2:].strip()
        # remove bold markers
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
        i += 1
        continue

    # Checkbox list
    if stripped.startswith("- [ ]") or stripped.startswith("- [x]"):
        checked = stripped.startswith("- [x]")
        text = stripped[5:].strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        marker = "[x] " if checked else "[ ] "
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(marker + text)
        i += 1
        continue

    # Numbered list
    m = re.match(r"^(\d+)\.\s+(.*)", stripped)
    if m:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2))
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)
        i += 1
        continue

    # Empty line
    if not stripped:
        i += 1
        continue

    # Normal paragraph — strip bold/italic markers for simplicity
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    p = doc.add_paragraph(text, style="Normal")
    i += 1

# flush any remaining table
if table_buffer:
    add_md_table(doc, parse_md_table(table_buffer))

doc.save(OUT)
print(f"Saved: {OUT}")
